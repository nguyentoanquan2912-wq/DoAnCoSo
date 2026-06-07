"""
Script format báo cáo ĐACS:
1. Thêm khung viền trang bìa
2. Tạo Mục lục tự động
3. Tạo Danh mục bảng
4. Tạo Danh mục hình ảnh
5. Đánh số trang (trang bìa không số, mục lục i/ii/iii, nội dung 1/2/3)
"""
import os
import sys
import io
import copy
import shutil
from datetime import datetime

# Fix encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, "Báo cáo ĐACS.docx")
OUTPUT_PATH = os.path.join(BASE_DIR, "Báo cáo ĐACS.docx")

# Backup first
BACKUP_PATH = os.path.join(BASE_DIR, f"Báo cáo ĐACS_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")


def add_page_border_to_section(section, color="000000", size="12", space="24", art=None):
    """Thêm khung viền cho trang trong section"""
    sectPr = section._sectPr
    
    # Remove existing borders if any
    existing = sectPr.find(qn('w:pgBorders'))
    if existing is not None:
        sectPr.remove(existing)
    
    # Create page borders element
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")} w:offsetFrom="page">'
        f'  <w:top w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'</w:pgBorders>'
    )
    
    # Insert before any existing child elements (after pgSz if exists)
    pgSz = sectPr.find(qn('w:pgSz'))
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.addnext(pgBorders)
    elif pgSz is not None:
        pgSz.addnext(pgBorders)
    else:
        sectPr.insert(0, pgBorders)


def add_double_page_border(section, outer_color="000000", outer_size="12", 
                            inner_color="000000", inner_size="6", space="24"):
    """Thêm khung viền đôi đẹp cho trang bìa"""
    sectPr = section._sectPr
    
    # Remove existing borders
    existing = sectPr.find(qn('w:pgBorders'))
    if existing is not None:
        sectPr.remove(existing)
    
    # Tạo khung viền đơn đậm nét đẹp cho trang bìa HUTECH
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")} w:offsetFrom="page">'
        f'  <w:top w:val="thinThickSmallGap" w:sz="24" w:space="{space}" w:color="{outer_color}"/>'
        f'  <w:left w:val="thinThickSmallGap" w:sz="24" w:space="{space}" w:color="{outer_color}"/>'
        f'  <w:bottom w:val="thickThinSmallGap" w:sz="24" w:space="{space}" w:color="{outer_color}"/>'
        f'  <w:right w:val="thickThinSmallGap" w:sz="24" w:space="{space}" w:color="{outer_color}"/>'
        f'</w:pgBorders>'
    )
    
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.addnext(pgBorders)
    else:
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.addnext(pgBorders)
        else:
            sectPr.insert(0, pgBorders)


def create_section_break(doc, paragraph, break_type='nextPage'):
    """Tạo section break sau paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    sectPr = parse_xml(f'<w:sectPr {nsdecls("w")}>'
                       f'<w:type w:val="{break_type}"/>'
                       f'</w:sectPr>')
    pPr.append(sectPr)
    return sectPr


def add_page_number_to_footer(section, number_format='decimal', start_num=1):
    """Thêm số trang vào footer của section
    number_format: 'decimal' (1,2,3), 'lowerRoman' (i,ii,iii), 'upperRoman' (I,II,III)
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Clear existing footer content
    for para in footer.paragraphs:
        para.clear()
    
    # Use first paragraph or create one
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()
    
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    run = para.add_run()
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # Add PAGE field code
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)
    
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._element.append(instrText)
    
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fldChar_end)
    
    # Set page number format in section properties
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}/>')
        sectPr.append(pgNumType)
    
    pgNumType.set(qn('w:fmt'), number_format)
    pgNumType.set(qn('w:start'), str(start_num))


def clear_footer(section):
    """Xóa footer (không hiển thị số trang)"""
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para.clear()


def insert_toc_field(doc, paragraph):
    """Chèn field TOC tự động vào paragraph"""
    run = paragraph.add_run()
    
    # Begin field
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)
    
    # Field instruction - TOC with heading levels 1-3
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u '
        f'</w:instrText>'
    )
    run._element.append(instrText)
    
    # Separate
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._element.append(fldChar_sep)
    
    # Placeholder text
    run2 = paragraph.add_run("Nhấn Ctrl+A rồi F9 để cập nhật mục lục")
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.color.rgb = RGBColor(128, 128, 128)
    run2.font.italic = True
    
    # End field
    run3 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar_end)


def insert_table_of_figures_field(doc, paragraph, caption_label="Hình"):
    """Chèn field danh mục hình ảnh/bảng"""
    run = paragraph.add_run()
    
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)
    
    # TOC field for figures/tables
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\h \\z \\c "{caption_label}" '
        f'</w:instrText>'
    )
    run._element.append(instrText)
    
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._element.append(fldChar_sep)
    
    run2 = paragraph.add_run(f"Nhấn Ctrl+A rồi F9 để cập nhật danh mục {caption_label.lower()}")
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.color.rgb = RGBColor(128, 128, 128)
    run2.font.italic = True
    
    run3 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar_end)


def find_paragraph_index(doc, text_match):
    """Tìm index paragraph chứa text"""
    for i, para in enumerate(doc.paragraphs):
        if text_match in para.text:
            return i
    return -1


def add_heading_paragraph(doc, text, level=1, font_size=16, bold=True, 
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, before_para_index=None):
    """Thêm heading paragraph"""
    if before_para_index is not None and before_para_index < len(doc.paragraphs):
        # Insert before specific paragraph
        ref_para = doc.paragraphs[before_para_index]
        new_para = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="Heading{level}"/></w:pPr></w:p>')
        ref_para._element.addprevious(new_para)
        
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_para, doc)
    else:
        para = doc.add_heading(text, level=level)
    
    para.alignment = alignment
    if para.runs:
        for run in para.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    return para


def build_manual_toc(doc):
    """Xây dựng mục lục thủ công dựa trên headings"""
    toc_entries = []
    
    # Mapping heading styles
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        text = para.text.strip()
        
        if not text:
            continue
            
        if style_name == 'Heading 1':
            # Skip MỤC LỤC, DANH MỤC heading themselves
            if text in ['MỤC LỤC', 'DANH MỤC HÌNH ẢNH', 'DANH MỤC BẢNG', '']:
                continue
            toc_entries.append({'level': 1, 'text': text})
        elif style_name == 'Heading 2':
            toc_entries.append({'level': 2, 'text': text})
        elif style_name == 'Heading 3':
            # Only include real headings (not long paragraphs styled as Heading 3)
            if len(text) < 100:
                toc_entries.append({'level': 3, 'text': text})
    
    return toc_entries


def create_toc_paragraph(doc, entry, parent_element):
    """Tạo một dòng trong mục lục"""
    level = entry['level']
    text = entry['text']
    
    para_xml = parse_xml(f'<w:p {nsdecls("w")}><w:pPr></w:pPr></w:p>')
    parent_element.append(para_xml)
    
    from docx.text.paragraph import Paragraph
    para = Paragraph(para_xml, doc)
    
    # Indent based on level
    indent_map = {1: 0, 2: Cm(1), 3: Cm(2)}
    indent = indent_map.get(level, 0)
    
    pPr = para._element.get_or_add_pPr()
    
    if indent:
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="{indent}"/>')
        pPr.append(ind)
    
    # Add tab stop with dot leader for page number
    tabs = parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        f'  <w:tab w:val="right" w:leader="dot" w:pos="9072"/>'
        f'</w:tabs>'
    )
    pPr.append(tabs)
    
    # Add text
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    
    if level == 1:
        run.font.size = Pt(13)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(12)
        run.font.bold = False
    else:
        run.font.size = Pt(11)
        run.font.bold = False
        run.font.italic = True
    
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    return para


def format_report():
    print("=" * 60)
    print("BẮT ĐẦU FORMAT BÁO CÁO ĐACS")
    print("=" * 60)
    
    # Backup
    print(f"\n1. Tạo backup: {os.path.basename(BACKUP_PATH)}")
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    
    # Open document
    print("2. Mở file báo cáo...")
    doc = Document(DOCX_PATH)
    
    # =========================================
    # BƯỚC 1: TẠO SECTION BREAKS
    # =========================================
    print("\n3. Tạo section breaks...")
    
    # Tìm các vị trí quan trọng
    cover_end_idx = -1  # Kết thúc trang bìa
    toc_idx = -1  # Mục lục
    content_start_idx = -1  # Bắt đầu nội dung chính
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ''
        
        if 'Thành phố Hồ Chí Minh' in text:
            cover_end_idx = i
        elif text == 'LỜI MỞ ĐẦU' and style == 'Heading 1':
            content_start_idx = i
        elif text == 'MỤC LỤC' and style == 'Heading 1':
            toc_idx = i
    
    print(f"   Trang bìa kết thúc tại P{cover_end_idx}")
    print(f"   Mục lục tại P{toc_idx}")  
    print(f"   Nội dung chính bắt đầu tại P{content_start_idx}")
    
    # =========================================
    # BƯỚC 2: THÊM KHUNG VIỀN TRANG BÌA
    # =========================================
    print("\n4. Thêm khung viền trang bìa...")
    
    # Cần tạo section break sau trang bìa để khung chỉ áp dụng cho trang bìa
    # Tìm paragraph cuối của trang bìa (trước LỜI MỞ ĐẦU)
    
    # Đầu tiên xóa page break cũ ở trang bìa nếu có và thay bằng section break
    # Trang bìa: từ P0 -> P13 (Thành phố HCM)
    
    # Tìm paragraph ngay trước LỜI MỞ ĐẦU  
    loi_mo_dau_idx = find_paragraph_index(doc, 'LỜI MỞ ĐẦU')
    loi_cam_doan_idx = find_paragraph_index(doc, 'LỜI CAM ĐOAN')
    muc_luc_idx = find_paragraph_index(doc, 'MỤC LỤC')
    
    # Tìm paragraph đầu tiên của CHƯƠNG 1
    chuong1_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if 'CHƯƠNG 1' in para.text and para.style and para.style.name == 'Heading 1':
            chuong1_idx = i
            break
    
    print(f"   LỜI MỞ ĐẦU: P{loi_mo_dau_idx}")
    print(f"   LỜI CAM ĐOAN: P{loi_cam_doan_idx}")
    print(f"   MỤC LỤC: P{muc_luc_idx}")
    print(f"   CHƯƠNG 1: P{chuong1_idx}")
    
    # =========================================
    # BƯỚC 3: Tạo sections bằng section breaks
    # =========================================
    # Section 1: Trang bìa (có khung viền, không số trang)
    # Section 2: Phần mở đầu (LỜI MỞ ĐẦU, LỜI CAM ĐOAN, MỤC LỤC, DANH MỤC) - số La Mã
    # Section 3: Nội dung chính (CHƯƠNG 1 trở đi) - số Ả Rập
    
    # Tạo section break sau trang bìa
    # Paragraph trước LỜI MỞ ĐẦU
    if loi_mo_dau_idx > 0:
        cover_last_para = doc.paragraphs[loi_mo_dau_idx - 1]
        # Nếu paragraph trước lời mở đầu trống, dùng nó
        # Nếu không, dùng paragraph kết thúc trang bìa
        if not cover_last_para.text.strip():
            target_para = cover_last_para
        else:
            target_para = doc.paragraphs[cover_end_idx] if cover_end_idx >= 0 else cover_last_para
        
        # Thêm section break
        pPr = target_para._element.get_or_add_pPr()
        # Remove existing section break if any
        existing_sectPr = pPr.find(qn('w:sectPr'))
        if existing_sectPr is not None:
            pPr.remove(existing_sectPr)
        
        # Copy section properties from first section
        first_sectPr = doc.sections[0]._sectPr
        new_sectPr = copy.deepcopy(first_sectPr)
        
        # Set as next page break
        type_elem = new_sectPr.find(qn('w:type'))
        if type_elem is None:
            type_elem = parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
            new_sectPr.insert(0, type_elem)
        else:
            type_elem.set(qn('w:val'), 'nextPage')
        
        pPr.append(new_sectPr)
        print("   Tạo section break sau trang bìa")
    
    # Tạo section break trước CHƯƠNG 1
    if chuong1_idx > 0:
        # Tìm paragraph trước CHƯƠNG 1
        pre_ch1_para = doc.paragraphs[chuong1_idx - 1]
        
        pPr = pre_ch1_para._element.get_or_add_pPr()
        existing_sectPr = pPr.find(qn('w:sectPr'))
        if existing_sectPr is not None:
            pPr.remove(existing_sectPr)
        
        # Copy section properties
        if len(doc.sections) > 0:
            base_sectPr = doc.sections[0]._sectPr
            new_sectPr = copy.deepcopy(base_sectPr)
        else:
            new_sectPr = parse_xml(f'<w:sectPr {nsdecls("w")}></w:sectPr>')
        
        type_elem = new_sectPr.find(qn('w:type'))
        if type_elem is None:
            type_elem = parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
            new_sectPr.insert(0, type_elem)
        else:
            type_elem.set(qn('w:val'), 'nextPage')
        
        # Remove page borders from this section (only cover has borders)
        borders = new_sectPr.find(qn('w:pgBorders'))
        if borders is not None:
            new_sectPr.remove(borders)
        
        pPr.append(new_sectPr)
        print("   Tạo section break trước CHƯƠNG 1")
    
    # =========================================
    # BƯỚC 4: THÊM KHUNG VIỀN CHO TRANG BÌA
    # =========================================
    print("\n5. Áp dụng khung viền cho trang bìa...")
    
    # Section đầu tiên = trang bìa
    if len(doc.sections) > 0:
        cover_section = doc.sections[0]
        add_double_page_border(cover_section, outer_color="000000")
        print("   Đã thêm khung viền đôi cho trang bìa")
    
    # =========================================
    # BƯỚC 5: XỬ LÝ MỤC LỤC  
    # =========================================
    print("\n6. Xử lý Mục lục...")
    
    # Tìm lại vị trí MỤC LỤC sau khi thêm sections
    muc_luc_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'MỤC LỤC' and para.style and 'Heading' in para.style.name:
            muc_luc_idx = i
            break
    
    if muc_luc_idx >= 0:
        muc_luc_para = doc.paragraphs[muc_luc_idx]
        
        # Xóa nội dung cũ sau MỤC LỤC (cho đến heading tiếp theo)
        # Tìm heading tiếp theo sau MỤC LỤC
        next_heading_idx = -1
        for i in range(muc_luc_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.style and 'Heading' in para.style.name and para.text.strip():
                next_heading_idx = i
                break
        
        # Xóa các paragraph giữa MỤC LỤC và heading tiếp theo
        if next_heading_idx > muc_luc_idx + 1:
            for i in range(next_heading_idx - 1, muc_luc_idx, -1):
                para = doc.paragraphs[i]
                if not para.text.strip() and not (para.style and 'Heading' in para.style.name):
                    para._element.getparent().remove(para._element)
        
        # Thêm TOC field sau heading MỤC LỤC
        # Tạo paragraph mới sau heading mục lục
        new_para_xml = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        muc_luc_para._element.addnext(new_para_xml)
        
        from docx.text.paragraph import Paragraph
        toc_para = Paragraph(new_para_xml, doc)
        
        # Insert TOC field
        insert_toc_field(doc, toc_para)
        print("   Đã chèn field Mục lục tự động")
    
    # =========================================
    # BƯỚC 6: XỬ LÝ DANH MỤC HÌNH ẢNH
    # =========================================
    print("\n7. Xử lý Danh mục hình ảnh...")
    
    dm_hinh_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'DANH MỤC HÌNH ẢNH' and para.style and 'Heading' in para.style.name:
            dm_hinh_idx = i
            break
    
    if dm_hinh_idx >= 0:
        dm_hinh_para = doc.paragraphs[dm_hinh_idx]
        
        # Tìm heading tiếp theo
        next_heading_idx = -1
        for i in range(dm_hinh_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.style and 'Heading' in para.style.name and para.text.strip():
                next_heading_idx = i
                break
        
        # Xóa paragraphs trống giữa DANH MỤC HÌNH ẢNH và heading tiếp
        if next_heading_idx > dm_hinh_idx + 1:
            for i in range(next_heading_idx - 1, dm_hinh_idx, -1):
                para = doc.paragraphs[i]
                if not (para.style and 'Heading' in para.style.name):
                    para._element.getparent().remove(para._element)
        
        # Tạo bảng danh mục hình ảnh thủ công
        # Tìm tất cả hình ảnh trong tài liệu
        figures = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith('Hình ') and ':' in text and len(text) < 120:
                figures.append(text)
        
        # Thêm nội dung danh mục
        if figures:
            for fig_text in figures:
                fig_para_xml = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
                dm_hinh_para._element.addnext(fig_para_xml)
                
                from docx.text.paragraph import Paragraph
                fig_para = Paragraph(fig_para_xml, doc)
                
                # Tab stop with dot leader
                pPr = fig_para._element.get_or_add_pPr()
                tabs = parse_xml(
                    f'<w:tabs {nsdecls("w")}>'
                    f'  <w:tab w:val="right" w:leader="dot" w:pos="9072"/>'
                    f'</w:tabs>'
                )
                pPr.append(tabs)
                
                run = fig_para.add_run(fig_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            print(f"   Đã thêm {len(figures)} mục hình ảnh")
        else:
            # Thêm placeholder
            placeholder_xml = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
            dm_hinh_para._element.addnext(placeholder_xml)
            from docx.text.paragraph import Paragraph
            ph_para = Paragraph(placeholder_xml, doc)
            run = ph_para.add_run("(Cập nhật sau khi hoàn thiện báo cáo)")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            print("   Không tìm thấy hình ảnh, thêm placeholder")
    else:
        print("   Không tìm thấy heading DANH MỤC HÌNH ẢNH")
    
    # =========================================
    # BƯỚC 7: THÊM DANH MỤC BẢNG
    # =========================================
    print("\n8. Thêm Danh mục bảng...")
    
    # Tìm vị trí DANH MỤC HÌNH ẢNH để chèn DANH MỤC BẢNG trước nó
    dm_hinh_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if 'DANH MỤC HÌNH ẢNH' in para.text and para.style and 'Heading' in para.style.name:
            dm_hinh_idx = i
            break
    
    if dm_hinh_idx >= 0:
        dm_hinh_para = doc.paragraphs[dm_hinh_idx]
        
        # Tạo heading DANH MỤC BẢNG trước DANH MỤC HÌNH ẢNH
        dm_bang_heading_xml = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'<w:pPr><w:pStyle w:val="Heading1"/><w:jc w:val="center"/></w:pPr>'
            f'</w:p>'
        )
        dm_hinh_para._element.addprevious(dm_bang_heading_xml)
        
        from docx.text.paragraph import Paragraph
        dm_bang_heading = Paragraph(dm_bang_heading_xml, doc)
        run = dm_bang_heading.add_run("DANH MỤC BẢNG")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Tìm tất cả bảng (dựa trên caption)
        tables_list = []
        table_names = [
            "Bảng 2.1: Ví dụ phân loại tin thật và tin giả",
            "Bảng 3.1: Danh sách API hệ thống",
            "Bảng 3.2: Cấu trúc bảng predictions trong SQLite",
            "Bảng 4.1: Công nghệ sử dụng trong hệ thống",
            "Bảng 4.2: Kết quả kiểm thử chức năng",
            "Bảng 4.3: Kết quả đánh giá mô hình Linear SVM",
        ]
        
        # Tìm bảng thực tế trong tài liệu
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith('Bảng ') and ':' in text and len(text) < 120:
                tables_list.append(text)
        
        if not tables_list:
            tables_list = table_names  # Dùng danh sách mặc định
        
        # Thêm các mục bảng
        # Chèn theo thứ tự ngược vì addnext sẽ đẩy các mục ra sau
        for tbl_text in reversed(tables_list):
            tbl_para_xml = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
            dm_bang_heading_xml.addnext(tbl_para_xml)
            
            tbl_para = Paragraph(tbl_para_xml, doc)
            
            pPr = tbl_para._element.get_or_add_pPr()
            tabs = parse_xml(
                f'<w:tabs {nsdecls("w")}>'
                f'  <w:tab w:val="right" w:leader="dot" w:pos="9072"/>'
                f'</w:tabs>'
            )
            pPr.append(tabs)
            
            run = tbl_para.add_run(tbl_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Thêm page break trước DANH MỤC HÌNH ẢNH
        pagebreak_xml = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'<w:r><w:br w:type="page"/></w:r>'
            f'</w:p>'
        )
        dm_hinh_para._element.addprevious(pagebreak_xml)
        
        print(f"   Đã thêm heading DANH MỤC BẢNG với {len(tables_list)} mục")
    
    # =========================================
    # BƯỚC 8: ĐÁNH SỐ TRANG
    # =========================================
    print("\n9. Đánh số trang...")
    
    # Refresh sections list
    sections = doc.sections
    print(f"   Tổng số sections: {len(sections)}")
    
    for i, section in enumerate(sections):
        print(f"   Section {i}: processing...")
    
    if len(sections) >= 3:
        # Section 0: Trang bìa - KHÔNG đánh số trang
        clear_footer(sections[0])
        print("   Section 0 (Trang bìa): Không đánh số trang")
        
        # Section 1: Phần mở đầu - Số La Mã (i, ii, iii)
        add_page_number_to_footer(sections[1], number_format='lowerRoman', start_num=1)
        print("   Section 1 (Phần mở đầu): Số La Mã (i, ii, iii...)")
        
        # Section 2+: Nội dung chính - Số Ả Rập (1, 2, 3)
        for i in range(2, len(sections)):
            add_page_number_to_footer(sections[i], number_format='decimal', start_num=1)
            print(f"   Section {i} (Nội dung): Số trang (1, 2, 3...)")
    elif len(sections) >= 2:
        # Nếu chỉ có 2 sections
        clear_footer(sections[0])
        print("   Section 0 (Trang bìa): Không đánh số trang")
        
        add_page_number_to_footer(sections[1], number_format='decimal', start_num=1)
        print("   Section 1 (Nội dung): Số trang (1, 2, 3...)")
    else:
        # Chỉ có 1 section
        add_page_number_to_footer(sections[0], number_format='decimal', start_num=1)
        print("   Section 0: Số trang (1, 2, 3...)")
    
    # =========================================
    # BƯỚC 9: XÓA PAGE BORDERS KHỎI CÁC SECTION KHÔNG PHẢI TRANG BÌA
    # =========================================
    print("\n10. Đảm bảo chỉ trang bìa có khung viền...")
    for i in range(1, len(sections)):
        borders = sections[i]._sectPr.find(qn('w:pgBorders'))
        if borders is not None:
            sections[i]._sectPr.remove(borders)
            print(f"    Đã xóa khung viền khỏi section {i}")
    
    # =========================================
    # BƯỚC 10: LƯU FILE
    # =========================================
    print(f"\n11. Lưu file...")
    doc.save(OUTPUT_PATH)
    print(f"    Đã lưu: {OUTPUT_PATH}")
    print(f"    Backup: {BACKUP_PATH}")
    
    print("\n" + "=" * 60)
    print("HOÀN TẤT FORMAT BÁO CÁO!")
    print("=" * 60)
    print("\nHƯỚNG DẪN:")
    print("1. Mở file báo cáo bằng WPS Office hoặc Microsoft Word")
    print("2. Nhấn Ctrl+A (chọn tất cả)")
    print("3. Nhấn F9 (cập nhật field)")
    print("4. Chọn 'Update entire table' nếu được hỏi")
    print("5. Kiểm tra và điều chỉnh nếu cần")


if __name__ == "__main__":
    format_report()
