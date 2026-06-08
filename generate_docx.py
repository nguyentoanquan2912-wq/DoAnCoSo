import sys
import os
import subprocess

# Tự động cài đặt python-docx nếu chưa có
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Thiết lập màu nền cho ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập padding cho ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()

    # Thiết lập lề trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Thống nhất font chữ mặc định là Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2B, 0x2D, 0x42) # Dark grey

    # --- Tiêu đề chính ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("BÁO CÁO THUYẾT TRÌNH CHI TIẾT\nPHẦN CÔNG NGHỆ AI & MÁY HỌC (TRUSTCHECK AI)")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81) # Classic Blue

    # --- Tiêu đề phụ ---
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    sub_run = subtitle_p.add_run("Tài liệu chuẩn bị cho buổi báo cáo với Giảng viên hướng dẫn")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D) # Grey

    # --- Đường kẻ ngang trang trí ---
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(24)
    run_line = p_line.add_run("━" * 45)
    run_line.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF) # Cyan

    # --- Danh sách các slide ---
    slides_data = [
        {
            "num": "Slide 1",
            "title": "TỔNG QUAN KIẾN TRÚC AI & NLP CỦA TRUSTCHECK",
            "bullets": [
                "Mô hình cốt lõi: Phân loại nhị phân sử dụng thuật toán học máy Linear SVM (Support Vector Machine).",
                "Hiệu suất thực nghiệm: Đạt độ chính xác (Accuracy) ~93 - 95% ổn định trên tập dữ liệu kiểm thử.",
                "Công nghệ xử lý tiếng Việt: Tích hợp thư viện underthesea để thực hiện tiền xử lý tách từ ghép tiếng Việt.",
                "Đặc điểm vận hành: Xử lý hoàn toàn cục bộ (local) trên máy chủ nội bộ, không phụ thuộc vào kết nối mạng bên ngoài và bảo mật thông tin tuyệt đối."
            ],
            "script": "Kính thưa Thầy/Cô và các bạn, trái tim của hệ thống TrustCheck AI là mô hình học máy Linear SVM kết hợp với quy trình xử lý ngôn ngữ tự nhiên tối ưu riêng cho tiếng Việt. Hệ thống đạt độ chính xác từ 93 đến 95% và chạy hoàn toàn cục bộ trên máy chủ nội bộ của chúng ta, giúp bảo vệ dữ liệu người dùng mà không cần phụ thuộc hay gửi dữ liệu cho bên thứ ba."
        },
        {
            "num": "Slide 2",
            "title": "QY TRÌNH TIỀN XỬ LÝ & TRÍCH XUẤT ĐẶC TRƯNG VĂN BẢN",
            "bullets": [
                "NLP tiếng Việt (underthesea): Tách từ ghép tiếng Việt (ví dụ: 'tin_tức', 'độ_tin_cậy') giúp mô hình phân tích ngữ nghĩa chính xác theo cụm từ.",
                "Tối ưu hóa trọng số Tiêu đề: Tự động nhân đôi tiêu đề bài viết (Tiêu đề + Tiêu đề + Nội dung) để các từ khóa ở tiêu đề có trọng số lớn hơn khi tính toán.",
                "Vector hóa TF-IDF nâng cao (ngram_range từ 1 đến 3): Trích xuất cả từ đơn (unigram), cụm 2 từ (bigram) và cụm 3 từ (trigram).",
                "Sublinear TF: Áp dụng công thức 1 + log(tf) giúp giảm bớt sự ảnh hưởng tiêu cực của các từ xuất hiện lặp lại quá nhiều lần trong một văn bản.",
                "Giới hạn đặc trưng: Cắt giảm tối đa 20,000 đặc trưng tối ưu nhất để tránh hiện tượng quá khớp (overfitting) và tối ưu hóa tốc độ."
            ],
            "script": "Để máy tính có thể hiểu được văn bản, chúng em xây dựng pipeline tiền xử lý qua 3 bước. Đầu tiên, chúng em sử dụng thư viện underthesea để tách từ ghép tiếng Việt. Tiếp theo, tiêu đề bài viết được nhân đôi trọng số để tăng độ ảnh hưởng trong TF-IDF. Cuối cùng, chúng em sử dụng bộ TF-IDF Vectorizer nâng cao trích xuất unigram, bigram và trigram lên đến 20,000 đặc trưng, áp dụng thuật toán sublinear scaling để loại bỏ độ nhiễu của tần suất từ."
        },
        {
            "num": "Slide 3",
            "title": "MÔ HÌNH LINEAR SVM & CƠ CHẾ HIỆU CHỈNH XÁC SUẤT",
            "bullets": [
                "Lựa chọn thuật toán: Linear SVM được lựa chọn vì hiệu năng xử lý văn bản tiếng Việt xuất sắc trên không gian đặc trưng thưa.",
                "Tham số tối ưu hóa: C=2.0 (biên phạt lỗi phân loại sai) và class_weight='balanced' để tự động xử lý hiện tượng mất cân bằng dữ liệu.",
                "Hiệu chỉnh xác suất (CalibratedClassifierCV): SVM truyền thống chỉ phân loại nhãn cứng dựa trên khoảng cách hình học tới siêu phẳng phân cách.",
                "Chúng em đã bọc SVM trong bộ hiệu chỉnh xác suất (CalibratedClassifierCV với 3-fold cross-validation) để chuyển đổi khoảng cách phân lớp thành điểm xác suất phần trăm (0 - 100%) hiển thị trực quan cho người dùng."
            ],
            "script": "Về phần mô hình, SVM tuyến tính là thuật toán có hiệu năng vượt trội nhất đối với phân loại văn bản. Tuy nhiên, điểm yếu của SVM nguyên bản là chỉ trả về nhãn cứng (Real/Fake) thay vì xác suất. Do đó, chúng em đã cải tiến bằng cách bọc SVM trong lớp CalibratedClassifierCV với cross-validation. Nhờ vậy, hệ thống có thể trả ra điểm số tin cậy cụ thể từ 0 đến 100%, giúp người dùng dễ dàng định lượng mức độ uy tín của tin tức."
        },
        {
            "num": "Slide 4",
            "title": "THUẬT TOÁN ĐỒNG THUẬN (CONSENSUS) & GIẢI THÍCH LÝ DO HEURISTICS",
            "bullets": [
                "Bỏ phiếu có trọng số (Weighted Voting): Tích hợp kết quả dự đoán của mô hình với các trọng số được thiết lập sẵn (Trọng số SVM là 2.0).",
                "Công thức tính Điểm Tin Cậy Tổng Hợp: Điểm Tin Cậy = (Xác suất trung bình * 0.6) + (Tỉ lệ đồng thuận * 100 * 0.4).",
                "Phân tích dấu hiệu Heuristics (NLP dựa trên quy luật):",
                "   - Giọng văn: Quét các từ giật gân, cường điệu hóa (ví dụ: 'sốc', 'rúng động', 'không thể tin nổi').",
                "   - Nguồn dẫn: Phát hiện sự tồn tại của từ chỉ nguồn trích dẫn pháp lý hoặc khoa học.",
                "   - Kêu gọi hành động khẩn cấp: Nhận diện các từ dẫn dụ chia sẻ gấp (ví dụ: 'chia sẻ ngay', 'share gấp')."
            ],
            "script": "Để kết quả phân tích có tính thuyết phục cao hơn, hệ thống TrustCheck không chỉ đưa ra điểm số từ mô hình máy học, mà còn tích hợp cơ chế phân tích dấu hiệu ngôn ngữ (Heuristics). Chúng em kiểm tra tính giật gân của câu chữ, sự hiện diện của các nguồn trích dẫn uy tín và các cụm từ kêu gọi hành động gấp. Kết quả cuối cùng là sự kết hợp có trọng số giữa dự đoán máy học và các quy luật ngôn ngữ học thực tế."
        },
        {
            "num": "Slide 5",
            "title": "TRÍCH XUẤT VĂN BẢN TỪ ẢNH CHỤP MÀN HÌNH (LOCAL OCR)",
            "bullets": [
                "Công nghệ: Sử dụng EasyOCR hoạt động hoàn toàn cục bộ, hỗ trợ nhận dạng chữ viết tiếng Việt và tiếng Anh.",
                "Tối ưu hóa tốc độ: Tự động resize ảnh về kích thước tối đa 1200px (giữ nguyên tỷ lệ) bằng bộ lọc LANCZOS giúp giảm 50% thời gian xử lý OCR.",
                "Thuật toán phân tích bố cục (Layout Heuristics):",
                "   - Tìm Tiêu đề: Quét các dòng chữ có chiều cao font chữ lớn nhất nằm ở 1/3 phía trên bức ảnh.",
                "   - Gom cụm đoạn văn: Đo khoảng cách chiều dọc (gap) giữa các dòng chữ. Nếu khoảng cách vượt quá 1.2 lần chiều cao font trung bình, hệ thống sẽ phân tách thành đoạn văn mới."
            ],
            "script": "Một tính năng rất thực tiễn của TrustCheck là phân tích bài báo qua ảnh chụp màn hình bằng EasyOCR. Để khắc phục tốc độ xử lý của OCR trên phần cứng thông thường, chúng em tự động thu nhỏ ảnh về ngưỡng 1200px sử dụng bộ lọc LANCZOS trước khi quét chữ. Đồng thời, chúng em thiết kế thuật toán phân tích layout: dòng chữ to nhất ở vùng trên cùng sẽ được tách làm Tiêu đề, các dòng chữ còn lại được gom nhóm thành các Đoạn văn dựa trên khoảng cách giữa các dòng."
        },
        {
            "num": "Slide 6",
            "title": "TRỢ LÝ CHATBOT AI NỘI BỘ VÀ QUẢN LÝ PHẠM VI HỘI THOẠI",
            "bullets": [
                "Bộ phân loại ý định (Intent Classifier): Sử dụng biểu thức chính quy (Regex) và phân tích cú pháp NLP để nhận biết nhanh 10+ ý định của người dùng.",
                "Hỗ trợ nghiệp vụ chuyên sâu: Trả lời về quy trình phân tích, giải thích điểm tin cậy, giải thích kết quả kiểm chứng, hướng dẫn dùng URL/OCR và cung cấp mẹo tự kiểm chứng nguồn tin.",
                "Quản lý ngoài phạm vi (Out of Scope): Nhận diện các câu hỏi lạc đề (nấu ăn, tài chính, bói toán, lập trình...) để lịch sự từ chối và định hướng người dùng quay lại chủ đề chính.",
                "Bảo mật: Hoạt động offline 100%, không gửi dữ liệu đoạn chat lên máy chủ Cloud bên ngoài."
            ],
            "script": "Cuối cùng là Trợ lý AI chatbot chuyên sâu. Chatbot này được tích hợp ngay trên giao diện nhằm giải đáp các câu hỏi thường gặp của người dùng về tin giả, hướng dẫn sử dụng tính năng cào link URL, OCR ảnh, đọc biểu đồ Dashboard hay tra cứu lịch sử. Nếu người dùng hỏi các câu hỏi không liên quan, chatbot sẽ lịch sự từ chối để tập trung hoàn toàn vào phạm vi kiểm chứng tin tức, mang lại trải nghiệm chuyên nghiệp và nhất quán."
        },
        {
            "num": "Slide 7",
            "title": "TỔNG HỢP CÁC CÔNG NGHỆ SỬ DỤNG TRONG TOÀN BỘ DỰ ÁN",
            "bullets": [
                "Python (Ngôn ngữ lập trình chính): Đóng vai trò là ngôn ngữ lập trình đồng bộ cho xử lý dữ liệu, máy học và backend Flask.",
                "Flask Framework: Thiết kế Web nhẹ và linh hoạt, được dùng để xây dựng các API Endpoints và render giao diện Jinja2.",
                "scikit-learn: Thư viện máy học dùng để huấn luyện mô hình Linear SVM, trích xuất đặc trưng TF-IDF, và hiệu chuẩn CalibratedClassifierCV.",
                "underthesea (NLP tiếng Việt): Thư viện xử lý ngôn ngữ tự nhiên tối ưu để tách từ ghép (word segmentation) và chuẩn hóa văn bản tiếng Việt.",
                "EasyOCR + Pillow (PIL): Thư viện Nhận dạng ký tự quang học (OCR) chạy cục bộ, trích xuất văn bản từ hình ảnh tin tức không cần API Key bên ngoài.",
                "SQLite (Cơ sở dữ liệu): Lưu trữ dữ liệu lịch sử thẩm định và dữ liệu huấn luyện cục bộ an toàn, gọn nhẹ.",
                "BeautifulSoup & Requests: Hỗ trợ cào dữ liệu (Web crawling) và trích xuất nội dung bài báo tự động từ đường link URL chính thống.",
                "HTML5 / Vanilla CSS / JavaScript: Giao diện người dùng hiện đại, responsive, kết hợp hoạt họa circuit board và phong cách Glassmorphism cao cấp."
            ],
            "script": "Để hoàn thiện hệ thống TrustCheck AI, chúng em đã sử dụng một hệ sinh thái công nghệ đồng bộ và mạnh mẽ. Trong đó, Python là ngôn ngữ nền tảng, kết hợp với Flask để vận hành backend, scikit-learn cho mô hình máy học SVM, underthesea để NLP tiếng Việt và EasyOCR cho OCR ảnh. Tất cả cơ sở dữ liệu được quản lý cục bộ bằng SQLite và giao diện được xây dựng hiện đại trên nền tảng HTML/CSS/JS thuần, đảm bảo tốc độ phản hồi tối ưu và khả năng chạy offline hoàn toàn."
        }
    ]

    for slide in slides_data:
        # Tiêu đề Slide
        h = doc.add_heading(level=2)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h_run = h.add_run(f"{slide['num']}: {slide['title']}")
        h_run.font.size = Pt(14)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81) # Classic Blue

        # Danh sách ý chính
        for bullet in slide['bullets']:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.add_run(bullet)

        # Khoảng cách trước hộp kịch bản nói
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Hộp chứa Kịch bản nói (Speech Script) dùng Table 1x1 để tạo khung trang trí
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        
        cell = table.cell(0, 0)
        set_cell_background(cell, "F4F6F9") # Màu nền xám nhạt sang trọng
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180) # Thêm padding
        
        # Bỏ viền bảng mặc định để tạo hiệu ứng blockquote hiện đại
        # (docx không hỗ trợ xóa viền trực tiếp một cách đơn giản, ta thiết kế kiểu màu xám nhạt bo góc mềm mại)

        cell_p = cell.paragraphs[0]
        cell_p.paragraph_format.space_after = Pt(0)
        run_label = cell_p.add_run("🎤 KỊCH BẢN NÓI THUYẾT TRÌNH (SPEAKER NOTES):\n")
        run_label.font.bold = True
        run_label.font.size = Pt(10.5)
        run_label.font.color.rgb = RGBColor(0x10, 0x85, 0xA8) # Teal Blue

        run_text = cell_p.add_run(f'"{slide["script"]}"')
        run_text.font.italic = True
        run_text.font.size = Pt(10.5)
        run_text.font.color.rgb = RGBColor(0x4A, 0x4E, 0x69) # Muted purple-grey

        # Khoảng cách sau mỗi slide
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Lưu file Word
    file_name = "Bao_Cao_AI_TrustCheck.docx"
    try:
        doc.save(file_name)
        # In bằng tiếng Anh đơn giản tránh lỗi Unicode console trên Windows
        print(f"SUCCESS: Saved file to {os.path.abspath(file_name)}")
    except PermissionError:
        file_name = "Bao_Cao_AI_TrustCheck_v2.docx"
        doc.save(file_name)
        print(f"FALLBACK: File was open. Saved to {os.path.abspath(file_name)}")

if __name__ == "__main__":
    create_report()
