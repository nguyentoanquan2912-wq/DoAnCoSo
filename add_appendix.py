"""
Script them Phu luc Hinh anh vao file bao cao .docx
Su dung python-docx de chen anh tu thu muc screenshots/
"""
import os
import sys
import io

# Fix encoding on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, "Báo cáo Đồ án Cơ sở.docx")
SCREENSHOT_DIR = os.path.join(BASE_DIR, "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "Báo cáo Đồ án Cơ sở.docx")  # Ghi đè file gốc

# Danh sách hình ảnh và mô tả
FIGURES = [
    {
        "file": "homepage.png",
        "caption": "Giao diện trang chủ TrustCheck AI",
        "desc": "Trang chủ hiển thị banner chính với tiêu đề \"Đánh giá độ tin cậy tin tức thông minh\", thanh điều hướng và các nút hành động chính."
    },
    {
        "file": "analyzer.png",
        "caption": "Giao diện trang Phân tích tin tức",
        "desc": "Trang phân tích cho phép người dùng nhập tiêu đề, nội dung bài báo hoặc URL để kiểm tra độ tin cậy. Hỗ trợ chọn mô hình ML và cung cấp mẫu nhanh."
    },
    {
        "file": "check_result.png",
        "caption": "Giao diện trang Kết quả phân tích",
        "desc": "Trang hiển thị kết quả phân tích chi tiết bao gồm: kết luận tin thật/giả, điểm tin cậy, phân tích lý do, từ khóa ảnh hưởng, biểu đồ trực quan và thẩm định chuyên sâu."
    },
    {
        "file": "dashboard.png",
        "caption": "Giao diện Dashboard hiệu suất mô hình ML",
        "desc": "Dashboard hiển thị các chỉ số đánh giá mô hình (Accuracy, Precision, Recall, F1-Score) và biểu đồ so sánh hiệu suất giữa các mô hình Machine Learning."
    },
    {
        "file": "history.png",
        "caption": "Giao diện trang Lịch sử dự đoán",
        "desc": "Trang lịch sử hiển thị danh sách các bài viết đã phân tích với thông tin thời gian, tiêu đề, kết quả, độ tin cậy và mô hình sử dụng. Hỗ trợ tìm kiếm, lọc và xóa."
    },
    {
        "file": "ai_assistant.png",
        "caption": "Giao diện Trợ lý AI Kiểm Chứng",
        "desc": "Trang trợ lý AI cho phép người dùng trò chuyện trực tiếp với hệ thống AI để kiểm chứng tin đồn, phân tích văn phong và nhận báo cáo thẩm định chuyên sâu."
    },
]


def add_appendix():
    print("Dang mo file bao cao...")
    doc = Document(DOCX_PATH)

    # Thêm ngắt trang trước phụ lục
    doc.add_page_break()

    # ---- TIÊU ĐỀ PHỤ LỤC ----
    heading = doc.add_heading("PHỤ LỤC HÌNH ẢNH", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Tùy chỉnh font heading
    for run in heading.runs:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Mô tả phụ lục
    desc_para = doc.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_para.add_run(
        "Phụ lục này trình bày các hình ảnh giao diện chính của hệ thống TrustCheck AI — "
        "Website Đánh Giá Độ Tin Cậy Tin Tức bằng Machine Learning và NLP."
    )
    desc_run.font.size = Pt(11)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # Khoảng trống

    # ---- DANH SÁCH HÌNH ẢNH ----
    figure_num = 1
    for fig in FIGURES:
        img_path = os.path.join(SCREENSHOT_DIR, fig["file"])
        if not os.path.exists(img_path):
            print(f"[SKIP] Khong tim thay: {fig['file']}")
            continue

        print(f"  Them hinh {figure_num}: {fig['caption']}")

        # Tiêu đề hình
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(f"Hình {figure_num}: {fig['caption']}")
        caption_run.font.size = Pt(11)
        caption_run.font.bold = True
        caption_run.font.color.rgb = RGBColor(0, 51, 102)

        # Chèn ảnh (width = 15cm để vừa trang A4)
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(img_path, width=Cm(15))

        # Mô tả hình
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.add_run(fig["desc"])
        desc_run.font.size = Pt(10)
        desc_run.font.italic = True
        desc_run.font.color.rgb = RGBColor(100, 100, 100)

        # Khoảng trống giữa các hình
        doc.add_paragraph()

        figure_num += 1

    # ---- BẢNG DANH MỤC HÌNH ẢNH ----
    doc.add_page_break()
    toc_heading = doc.add_heading("DANH MỤC HÌNH ẢNH", level=2)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_heading.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    # Tạo bảng danh mục
    table = doc.add_table(rows=1, cols=3)
    # Apply border to table using XML
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '000000',
        })
        borders.append(border)
    tblPr.append(borders)

    # Header
    hdr_cells = table.rows[0].cells
    headers = ["STT", "Hình ảnh", "Mô tả"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Tô nền header
        from docx.oxml.ns import qn
        shading_elm = hdr_cells[i]._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading_elm.append(shading)

    # Dữ liệu
    fig_num = 1
    for fig in FIGURES:
        img_path = os.path.join(SCREENSHOT_DIR, fig["file"])
        if not os.path.exists(img_path):
            continue

        row_cells = table.add_row().cells
        row_cells[0].text = str(fig_num)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = f"Hình {fig_num}"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].text = fig["caption"]

        # Font size cho các ô
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        fig_num += 1

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(11)

    # Lưu file
    doc.save(OUTPUT_PATH)
    print(f"\nDa luu file bao cao: {OUTPUT_PATH}")
    print(f"Tong so hinh anh da them: {figure_num - 1}")


if __name__ == "__main__":
    add_appendix()
